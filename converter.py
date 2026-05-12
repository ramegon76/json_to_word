#!/usr/bin/env python3
"""
SIM JSON → Word converter
Uso: python sim_json_to_word.py input.json output.docx

Genera un documento Word con formato institucional INEGI
a partir del JSON exportado del Sistema Integrador de Metadatos (SIM).

Dependencia:
    pip install python-docx
"""

import json
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colores corporativos ─────────────────────────────────────────────────────
AZUL  = RGBColor(0x00, 0x30, 0x57)
GRIS  = RGBColor(0x33, 0x33, 0x33)
BORDE = "E5E7EB"
FUENTE = "Arial"


# ─── Helpers XML ─────────────────────────────────────────────────────────────

def set_table_width(table, width_dxa=9360):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')


def write_cell(cell, text, bold=False, header=False, size=12):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(str(text or ''))
    run.font.name = FUENTE
    run.font.size = Pt(size)
    run.font.bold = bold or header
    run.font.color.rgb = AZUL if header else GRIS
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for tag in ('w:tcBorders', 'w:tcMar', 'w:shd', 'w:vAlign'):
        el = tcPr.find(qn(tag))
        if el is not None:
            tcPr.remove(el)

    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), BORDE)
        tcBorders.append(el)
    tcPr.append(tcBorders)

    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', 80), ('start', 160), ('bottom', 80), ('end', 160)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


# ─── Helpers de párrafo ───────────────────────────────────────────────────────

def add_para(doc, text='', bold=False, color=None, size=12,
             before=3, after=3, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = Pt(13.8)
    if keep:
        p.paragraph_format.keep_with_next = True
    if text:
        lines = str(text).split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.name = FUENTE
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color or GRIS
    return p


def add_spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph('')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


# ─── Estilos semánticos ───────────────────────────────────────────────────────

def add_chapter(doc, text):
    """Capítulo: Metadatos referenciales / estructurales / descriptivos"""
    p = add_para(doc, text, bold=False, color=AZUL, size=12, before=14, after=6)
    return p


def add_section(doc, text):
    """Sección: Identificación / Proceso de Producción / etc."""
    p = add_para(doc, text, bold=False, color=AZUL, size=12, before=10, after=4)
    return p


def add_subsection(doc, text):
    """Sub-sección: Creación o actualización / Datos de identificación / etc."""
    p = add_para(doc, text, bold=False, color=AZUL, size=12, before=8, after=3)
    return p


def add_label(doc, text):
    """Etiqueta de campo: negrita azul 11pt"""
    p = add_para(doc, text, bold=True, color=AZUL, size=12, before=6, after=2, keep=True)
    return p


def add_value(doc, text):
    """Valor de campo: gris normal 11pt"""
    p = add_para(doc, text, color=GRIS, size=12, before=0, after=4)
    return p


# ─── Tablas ───────────────────────────────────────────────────────────────────

def add_table(doc, headers, rows, after_spacer=True):
    if not rows:
        return
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    set_table_width(table, 9360)

    for j, h in enumerate(headers):
        write_cell(table.rows[0].cells[j], h, header=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            write_cell(table.rows[i + 1].cells[j], val or '')

    if after_spacer:
        add_para(doc, '', before=0, after=2)


def add_single_col_table(doc, values, after_spacer=True):
    values = [v for v in values if v and str(v).strip()]
    if not values:
        return
    table = doc.add_table(rows=len(values), cols=1)
    table.style = 'Table Grid'
    set_table_width(table, 9360)
    for i, val in enumerate(values):
        write_cell(table.rows[i].cells[0], val)
    if after_spacer:
        add_para(doc, '', before=0, after=2)


def add_fixed_2col_table(doc, headers, rows, col1_cm=8):
    """Tabla de 2 columnas con ancho fijo en col1 y resto para col2.
    Usa layout fijo para que URLs largas no expandan la columna."""
    if not rows:
        return
    # En landscape con márgenes 1.27cm: 11" = 15840 DXA - 2×(1.27cm×567) = 14400 DXA útiles
    total_dxa = 14400
    col1_dxa = int(col1_cm * 567)
    col2_dxa = total_dxa - col1_dxa
    widths = [col1_dxa, col2_dxa]

    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = 'Table Grid'
    set_table_width(table, total_dxa)

    # Layout fijo para respetar anchos sin importar el contenido
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    tblLayout = tblPr.find(qn('w:tblLayout')) if tblPr is not None else None
    if tblLayout is None and tblPr is not None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    if tblLayout is not None:
        tblLayout.set(qn('w:type'), 'fixed')

    # Grid
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(1, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)

    def set_col_width(cell, w):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = OxmlElement('w:tcW')
            tcPr.insert(0, tcW)
        tcW.set(qn('w:w'), str(w))
        tcW.set(qn('w:type'), 'dxa')

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_col_width(cell, widths[j])
        write_cell(cell, h, header=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            set_col_width(cell, widths[j])
            write_cell(cell, val or '')

    add_para(doc, '', before=0, after=2)


def render_list_or_text(doc, values):
    """Si hay más de un ítem → tabla de una columna. Si hay uno → texto plano."""
    values = [v for v in values if v and str(v).strip()]
    if not values:
        return
    if len(values) > 1:
        add_single_col_table(doc, values)
    else:
        add_value(doc, values[0])


# ─── Limpieza de valores ──────────────────────────────────────────────────────

def clean_cat(val):
    """Elimina prefijo de catálogo: '08 / Anual' o 'M_06 / Encuesta' → 'Anual' / 'Encuesta'"""
    if not val:
        return ''
    return re.sub(r'^[\w_]+\s*/\s*', '', str(val)).strip()


def remove_acronym(val):
    """Elimina acrónimo inicial: 'UGMA / Unidad de Geografía' → 'Unidad de Geografía'"""
    if not val:
        return ''
    parts = str(val).split(' / ', 1)
    return parts[1].strip() if len(parts) == 2 else str(val).strip()

def bool_str(val):
    return 'Sí' if str(val) in ('1', 'True', 'true') else 'No'


def fix_date(val):
    """Convierte YYYY-MM-DD → DD/MM/AAAA en cualquier string que contenga ese patrón."""
    if not val:
        return val
    return re.sub(
        r'(\d{4})-(\d{2})-(\d{2})',
        lambda m: f'{m.group(3)}/{m.group(2)}/{m.group(1)}',
        str(val)
    )


def fmt(elem, val):
    if val is None:
        return ''
    tipo = elem.get('claveTipoDato', '') or ''
    if tipo == 'EC':
        return clean_cat(val)
    if tipo == 'BO':
        return bool_str(val)
    # Convertir fechas en formato anglosajón al formato DD/MM/AAAA
    return fix_date(str(val).strip())


# ─── Extracción de datos ──────────────────────────────────────────────────────

def get_values(elem):
    """Retorna lista de valores simples ordenados."""
    return [
        fmt(elem, c['valor'])
        for c in sorted(elem.get('contenido', []), key=lambda x: x.get('orden', 0))
        if c.get('valor') is not None and str(c.get('valor', '')).strip()
    ]


def get_single(elem):
    vals = get_values(elem)
    return vals[0] if vals else ''


def get_compound_rows(elem):
    """Retorna (subfields, headers, rows) para elementos compuestos."""
    subfields = elem.get('elemento', [])
    if not subfields:
        return [], [], []
    parent_entries = sorted(
        [c for c in elem.get('contenido', []) if c.get('idPadreContenido') is None],
        key=lambda c: c.get('orden', 0)
    )
    parent_ids = [c['idContenido'] for c in parent_entries]
    if not parent_ids:
        return subfields, [sf['nombre'].strip() for sf in subfields], []

    headers = [sf['nombre'].strip() for sf in subfields]
    rows = []
    for pid in parent_ids:
        row = []
        for sf in subfields:
            match = next((c for c in sf.get('contenido', []) if c.get('idPadreContenido') == pid), None)
            val = match['valor'] if match else None
            row.append(fmt(sf, val) or '')
        if any(v for v in row):
            rows.append(row)
    return subfields, headers, rows


def has_content(elem):
    vals = [c for c in elem.get('contenido', []) if c.get('valor') is not None and str(c.get('valor','')).strip()]
    if vals:
        return True
    for sf in elem.get('elemento', []):
        if any(c.get('valor') is not None and str(c.get('valor','')).strip() for c in sf.get('contenido', [])):
            return True
    return False


def find_elem(seccion, nombre):
    """Busca un elemento por nombre exacto o parcial en una sección."""
    for e in seccion.get('elemento', []):
        if e['nombre'].strip() == nombre or nombre in e['nombre'].strip():
            return e
    return None


def find_subsec(seccion, nombre):
    for s in seccion.get('seccion', []):
        if nombre.lower() in s['nombre'].strip().lower():
            return s
    return None


def find_top_sec(data, nombre):
    for s in data.get('seccion', []):
        if nombre.lower() in s['nombre'].strip().lower():
            return s
    return None


# ─── Constructor del documento ────────────────────────────────────────────────

def build_document(data):
    doc = Document()

    # Configurar página carta con márgenes 0.75"
    sec = doc.sections[0]
    # Horizontal (landscape): ancho y alto se intercambian
    sec.page_width  = Inches(11)
    sec.page_height = Inches(8.5)
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, attr, Cm(1.27))
    # Indicar orientación landscape en el XML de la sección
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn
    pgSz = sec._sectPr.find(_qn('w:pgSz'))
    if pgSz is None:
        pgSz = _OxmlElement('w:pgSz')
        sec._sectPr.append(pgSz)
    pgSz.set(_qn('w:orient'), 'landscape')

    # ── ENCABEZADO ──────────────────────────────────────────────────────────
    nombre       = data.get('nombreInstancia', '')
    complemento  = data.get('complementoNombre', '')
    nombre_corto = data.get('nombreCortoInstancia', '')

    add_para(doc, nombre, bold=True, color=AZUL, size=12, before=0, after=4)
    if complemento:
        add_para(doc, complemento, bold=True, color=AZUL, size=12, before=0, after=4)
    add_para(doc, nombre_corto, bold=True, color=AZUL, size=12, before=0, after=0)
    add_spacer(doc, 1)

    # ── METADATOS REFERENCIALES ──────────────────────────────────────────────
    sec_ref = find_top_sec(data, 'referenciales')
    if sec_ref:
        add_chapter(doc, 'Metadatos referenciales')
        add_spacer(doc, 1)

        # — Identificación del documento de metadatos —
        sec_id = find_subsec(sec_ref, 'identificación')
        if sec_id:
            add_section(doc, 'Identificación del documento de metadatos')
            sec_creacion = find_subsec(sec_id, 'creación')
            if sec_creacion:
                add_subsection(doc, 'Creación o actualización del documento de metadatos')

                # Productor del metadato → tabla Nombre | Siglas | Dependencia
                e = find_elem(sec_creacion, 'Productor del metadato')
                if e and has_content(e):
                    add_label(doc, 'Productor del metadato')
                    _, hdrs, rows = get_compound_rows(e)
                    add_table(doc, hdrs, rows)

                # Fecha de actualización
                e = find_elem(sec_creacion, 'Fecha de actualización')
                if e and has_content(e):
                    add_label(doc, 'Fecha de actualización del metadato')
                    add_value(doc, get_single(e))

                # Actualización de los metadatos
                e = find_elem(sec_creacion, 'Actualización de los metadatos')
                if e and has_content(e):
                    add_label(doc, 'Actualización de los metadatos')
                    add_value(doc, get_single(e))

        add_spacer(doc, 1)

        # — Proceso de Producción documentado —
        sec_pp = find_subsec(sec_ref, 'producción documentado')
        if sec_pp:
            add_section(doc, 'Proceso de Producción documentado')
            sec_datos = find_subsec(sec_pp, 'datos de identificación')
            if sec_datos:
                add_subsection(doc, 'Datos de identificación')

                # Unidad Administrativa (sin acrónimo)
                e = find_elem(sec_datos, 'Unidad Administrativa')
                if e and has_content(e):
                    add_label(doc, 'Unidad Administrativa')
                    add_value(doc, remove_acronym(get_single(e)))

                # Proceso de Producción (sin acrónimo)
                e = find_elem(sec_datos, 'Proceso de Producción')
                if e and has_content(e):
                    add_label(doc, 'Proceso de Producción')
                    add_value(doc, remove_acronym(get_single(e)))

                # Programa de Información → columna única
                e = find_elem(sec_datos, 'Programa de Información')
                if e and has_content(e):
                    add_label(doc, 'Programa de Información')
                    render_list_or_text(doc, [clean_cat(v) for v in get_values(e)])

                # Ciclo de proceso
                e = find_elem(sec_datos, 'Ciclo de proceso')
                if e and has_content(e):
                    add_label(doc, 'Ciclo de proceso')
                    add_value(doc, get_single(e))

                # Subtítulo del documento de metadatos
                e = find_elem(sec_datos, 'Subtítulo')
                if e and has_content(e):
                    add_label(doc, 'Subtítulo del documento de metadatos')
                    add_value(doc, get_single(e))

        add_spacer(doc, 1)

        # — Características del Proceso de Producción —
        sec_car = find_subsec(sec_ref, 'características')
        if sec_car:
            add_section(doc, 'Características del Proceso de Producción')

            # Productores de los Procesos de Producción
            sec_prod = find_subsec(sec_car, 'productores')
            if sec_prod:
                add_subsection(doc, 'Productores de los Procesos de Producción de Información')

                # Área responsable → tabla Nombre | Siglas | Dependencia
                e = find_elem(sec_prod, 'Área responsable')
                if e and has_content(e):
                    add_label(doc, 'Área responsable del Proceso de Producción')
                    _, hdrs, rows = get_compound_rows(e)
                    add_table(doc, hdrs, rows)

                # Colaboradores → tabla Nombre | Siglas | Dependencia | Rol
                e = find_elem(sec_prod, 'Colaboradores')
                if e and has_content(e):
                    add_label(doc, 'Colaboradores')
                    _, hdrs, rows = get_compound_rows(e)
                    add_table(doc, hdrs, rows)

                # Financiamiento
                e = find_elem(sec_prod, 'Financiamiento')
                if e and has_content(e):
                    add_label(doc, 'Financiamiento')
                    add_value(doc, clean_cat(get_single(e)))

            # Descripción general
            sec_desc = find_subsec(sec_car, 'descripción general')
            if sec_desc:
                add_subsection(doc, 'Descripción general')

                e = find_elem(sec_desc, 'Objetivo')
                if e and has_content(e):
                    add_label(doc, 'Objetivo')
                    add_value(doc, get_single(e))

                e = find_elem(sec_desc, 'Antecedentes')
                if e and has_content(e):
                    add_label(doc, 'Antecedentes')
                    add_value(doc, get_single(e))

                e = find_elem(sec_desc, 'Método de generación')
                if e and has_content(e):
                    add_label(doc, 'Método de generación de la información')
                    add_value(doc, clean_cat(get_single(e)))

                e = find_elem(sec_desc, 'Periodicidad')
                if e and has_content(e):
                    add_label(doc, 'Periodicidad de producción')
                    render_list_or_text(doc, [clean_cat(v) for v in get_values(e)])

                e = find_elem(sec_desc, 'Grado de madurez')
                if e and has_content(e):
                    add_label(doc, 'Grado de madurez')
                    add_value(doc, clean_cat(get_single(e)))

                e = find_elem(sec_desc, 'Condición de Información de Interés Nacional')
                if e and has_content(e):
                    add_label(doc, 'Condición de Información de Interés Nacional')
                    add_value(doc, get_single(e))

                e = find_elem(sec_desc, 'Palabras clave')
                if e and has_content(e):
                    add_label(doc, 'Palabras clave')
                    add_value(doc, get_single(e))

                # Periodo de ejecución → tabla Inicio | Fin | Fase
                e = find_elem(sec_desc, 'Periodo de ejecución')
                if e and has_content(e):
                    add_label(doc, 'Periodo de ejecución')
                    _, hdrs, rows = get_compound_rows(e)
                    add_table(doc, hdrs, rows)

            # Diseño conceptual
            sec_dc = find_subsec(sec_car, 'diseño conceptual')
            if sec_dc:
                add_subsection(doc, 'Diseño conceptual')

                e = find_elem(sec_dc, 'Población objeto de estudio')
                if e and has_content(e):
                    add_label(doc, 'Población objeto de estudio')
                    add_value(doc, get_single(e))

                e = find_elem(sec_dc, 'Unidad de observación')
                if e and has_content(e):
                    add_label(doc, 'Unidad de observación o de análisis')
                    add_value(doc, get_single(e))

                e = find_elem(sec_dc, 'Indicadores objetivo')
                if e and has_content(e):
                    add_label(doc, 'Indicadores objetivo')
                    render_list_or_text(doc, get_values(e))

                e = find_elem(sec_dc, 'Dominios de estudio')
                if e and has_content(e):
                    add_label(doc, 'Dominios de estudio')
                    render_list_or_text(doc, get_values(e))

                e = find_elem(sec_dc, 'Cobertura temática')
                if e and has_content(e):
                    add_label(doc, 'Cobertura temática')
                    render_list_or_text(doc, get_values(e))

                e = find_elem(sec_dc, 'Representación espacial')
                if e and has_content(e):
                    add_label(doc, 'Representación espacial')
                    render_list_or_text(doc, [clean_cat(v) for v in get_values(e)])

                e = find_elem(sec_dc, 'Estándares y clasificaciones')
                if e and has_content(e):
                    add_label(doc, 'Estándares y clasificaciones')
                    render_list_or_text(doc, get_values(e))

                e = find_elem(sec_dc, 'Cobertura geográfica')
                if e and has_content(e):
                    add_label(doc, 'Cobertura geográfica')
                    add_value(doc, get_single(e))

            # Diseño muestral
            sec_dm = find_subsec(sec_car, 'diseño muestral')
            if sec_dm:
                elems_con_contenido = [e for e in sec_dm.get('elemento', []) if has_content(e)]
                if elems_con_contenido:
                    add_subsection(doc, 'Diseño muestral')

                    e = find_elem(sec_dm, 'Unidad de muestreo')
                    if e and has_content(e):
                        add_label(doc, 'Unidad de muestreo')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_dm, 'Diseño de la muestra')
                    if e and has_content(e):
                        add_label(doc, 'Diseño de la muestra')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_dm, 'Tamaño de la muestra')
                    if e and has_content(e):
                        add_label(doc, 'Tamaño de la muestra')
                        add_value(doc, get_single(e))

            # Captación de los datos
            sec_cap = find_subsec(sec_car, 'captación')
            if sec_cap:
                elems_con_contenido = [e for e in sec_cap.get('elemento', []) if has_content(e)]
                if elems_con_contenido:
                    add_subsection(doc, 'Captación de los datos')

                    # Periodo de referencia → tabla Inicio | Fin | Ciclo o Tema
                    e = find_elem(sec_cap, 'Periodo de referencia')
                    if e and has_content(e):
                        add_label(doc, 'Periodo de referencia')
                        _, hdrs, rows = get_compound_rows(e)
                        add_table(doc, hdrs, rows)

                    e = find_elem(sec_cap, 'Modalidad de captación')
                    if e and has_content(e):
                        add_label(doc, 'Modalidad de captación o recolección de los datos')
                        render_list_or_text(doc, [clean_cat(v) for v in get_values(e)])

                    e = find_elem(sec_cap, 'Responsable de la captación')
                    if e and has_content(e):
                        add_label(doc, 'Responsable de la captación de los datos')
                        render_list_or_text(doc, get_values(e))

                    e = find_elem(sec_cap, 'Descripción del instrumento')
                    if e and has_content(e):
                        add_label(doc, 'Descripción del instrumento de captación')
                        add_value(doc, get_single(e))

            # Procesamiento
            sec_proc = find_subsec(sec_car, 'procesamiento')
            if sec_proc:
                elems_con_contenido = [e for e in sec_proc.get('elemento', []) if has_content(e)]
                if elems_con_contenido:
                    add_subsection(doc, 'Procesamiento')

                    e = find_elem(sec_proc, 'Compilación de datos')
                    if e and has_content(e):
                        add_label(doc, 'Compilación de datos')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_proc, 'Procesamiento de datos')
                    if e and has_content(e):
                        add_label(doc, 'Procesamiento de datos')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_proc, 'Validación de datos')
                    if e and has_content(e):
                        add_label(doc, 'Validación de datos')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_proc, 'Imputación')
                    if e and has_content(e):
                        add_label(doc, 'Imputación')
                        add_value(doc, get_single(e))

            # Análisis
            sec_anal = find_subsec(sec_car, 'análisis')
            if sec_anal:
                elems_con_contenido = [e for e in sec_anal.get('elemento', []) if has_content(e)]
                if elems_con_contenido:
                    add_subsection(doc, 'Análisis')

                    e = find_elem(sec_anal, 'Análisis de datos')
                    if e and has_content(e):
                        add_label(doc, 'Análisis de datos')
                        add_value(doc, get_single(e))

            # Difusión
            sec_dif = find_subsec(sec_car, 'difusión')
            if sec_dif:
                elems_con_contenido = [e for e in sec_dif.get('elemento', []) if has_content(e)]
                if elems_con_contenido:
                    add_subsection(doc, 'Difusión')

                    # Resultados disponibles → tabla Nombre | URI
                    e = find_elem(sec_dif, 'Resultados disponibles')
                    if e and has_content(e):
                        add_label(doc, 'Resultados disponibles')
                        _, hdrs, rows = get_compound_rows(e)
                        add_table(doc, hdrs, rows)

                    e = find_elem(sec_dif, 'Estatus de la información')
                    if e and has_content(e):
                        add_label(doc, 'Estatus de la información')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_dif, 'Condición de acceso a microdatos')
                    if e and has_content(e):
                        add_label(doc, 'Condición de acceso a microdatos')
                        add_value(doc, clean_cat(get_single(e)))

                    e = find_elem(sec_dif, 'Política de confidencialidad')
                    if e and has_content(e):
                        add_label(doc, 'Política de confidencialidad')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_dif, 'Contacto')
                    if e and has_content(e):
                        add_label(doc, 'Contacto')
                        add_value(doc, get_single(e))

            # Aseguramiento de la calidad
            sec_cal = find_subsec(sec_car, 'aseguramiento')
            if sec_cal:
                elems_con_contenido = [e for e in sec_cal.get('elemento', []) if has_content(e)]
                if elems_con_contenido:
                    add_subsection(doc, 'Aseguramiento de la calidad')

                    e = find_elem(sec_cal, 'Evaluaciones implementadas')
                    if e and has_content(e):
                        add_label(doc, 'Evaluaciones implementadas')
                        render_list_or_text(doc, get_values(e))

                    # Pertinencia → tabla Fuente | Referencia
                    e = find_elem(sec_cal, 'Pertinencia')
                    if e and has_content(e):
                        add_label(doc, 'Pertinencia: Uso de la información')
                        _, hdrs, rows = get_compound_rows(e)
                        add_table(doc, hdrs, rows)

                    # Accesibilidad → tabla Nombre | URI
                    e = find_elem(sec_cal, 'Accesibilidad')
                    if e and has_content(e):
                        add_label(doc, 'Accesibilidad: Interpretación de la información')
                        _, hdrs, rows = get_compound_rows(e)
                        add_fixed_2col_table(doc, hdrs, rows, col1_cm=8)

                    e = find_elem(sec_cal, 'Puntualidad: Incorporación')
                    if e and has_content(e):
                        add_label(doc, 'Puntualidad: Incorporación en calendario de difusión')
                        add_value(doc, clean_cat(get_single(e)))

                    # Puntualidad: Fechas → tabla 3 columnas
                    e = find_elem(sec_cal, 'Puntualidad: Fechas')
                    if e and has_content(e):
                        add_label(doc, 'Puntualidad: Fechas de publicación comprometidas')
                        _, hdrs, rows = get_compound_rows(e)
                        add_table(doc, hdrs, rows)

                    e = find_elem(sec_cal, 'Oportunidad')
                    if e and has_content(e):
                        add_label(doc, 'Oportunidad: Parámetro de referencia y oportunidad observada')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_cal, 'Coherencia y comparabilidad')
                    if e and has_content(e):
                        add_label(doc, 'Coherencia y comparabilidad: Comparabilidad temporal y geográfica')
                        add_value(doc, get_single(e))

                    e = find_elem(sec_cal, 'Veracidad: Precisión')
                    if e and has_content(e):
                        add_label(doc, 'Veracidad: Precisión y confiabilidad estadística y exactitud geográfica')
                        add_value(doc, get_single(e))

                    # Veracidad: Indicadores → tabla Nombre indicador | URI
                    e = find_elem(sec_cal, 'Veracidad: Indicadores')
                    if e and has_content(e):
                        add_label(doc, 'Veracidad: Indicadores de precisión y confiabilidad estadística y exactitud geográfica')
                        _, hdrs, rows = get_compound_rows(e)
                        add_fixed_2col_table(doc, hdrs, rows, col1_cm=8)

    add_spacer(doc, 1)

    # ── METADATOS ESTRUCTURALES ──────────────────────────────────────────────
    sec_est = find_top_sec(data, 'estructurales')
    if sec_est:
        e_bds = find_elem(sec_est, 'Descripción del archivo')
        if e_bds and has_content(e_bds):
            add_chapter(doc, 'Metadatos estructurales')
            add_section(doc, 'Descripción del archivo de base de datos')

            subfields = e_bds.get('elemento', [])

            # El subfield "Archivo de base de datos" tiene un contenido por BD
            # con idPadreContenido apuntando al grupo raíz de cada BD.
            # Los IDs raíz (sin padre) son los grupos de cada base de datos.
            # Obtenemos los pid_bd como los idPadreContenido únicos que apuntan
            # a contenidos con idPadreContenido=None en el elemento principal.
            raiz_ids = sorted(
                [c['idContenido'] for c in e_bds.get('contenido', [])
                 if c.get('idPadreContenido') is None],
                key=lambda x: x
            )

            for pid_bd in raiz_ids:
                def get_sf_val(nombre_sf, pid=pid_bd):
                    sf = next((s for s in subfields
                                if nombre_sf.lower() in s['nombre'].strip().lower()), None)
                    if not sf:
                        return ''
                    match = next((c for c in sf.get('contenido', [])
                                  if c.get('idPadreContenido') == pid), None)
                    if match and match.get('valor') is not None and str(match['valor']).strip():
                        return fmt(sf, match['valor'])
                    return ''

                def get_sf_vals(nombre_sf, pid=pid_bd):
                    """Para campos multi-valor como Fuente de información."""
                    sf = next((s for s in subfields
                                if nombre_sf.lower() in s['nombre'].strip().lower()), None)
                    if not sf:
                        return []
                    return [
                        fmt(sf, c['valor'])
                        for c in sorted(sf.get('contenido', []), key=lambda x: x.get('orden', 0))
                        if c.get('idPadreContenido') == pid
                        and c.get('valor') is not None
                        and str(c['valor']).strip()
                    ]

                # Verificar que la BD tiene al menos un campo con contenido
                nombre_bd = get_sf_val('Nombre del archivo')
                if not nombre_bd:
                    continue

                add_label(doc, 'Nombre del archivo de base de datos')
                add_value(doc, nombre_bd)

                v = get_sf_val('Unidad y Área Administrativa')
                if v:
                    add_label(doc, 'Unidad y Área Administrativa responsable')
                    add_value(doc, v)

                v = get_sf_val('Sitio de descarga')
                if v:
                    add_label(doc, 'Sitio de descarga de la base de datos')
                    add_value(doc, v)

                v = get_sf_val('Contenido de la base de datos')
                if v:
                    add_label(doc, 'Contenido de la base de datos')
                    add_value(doc, v)

                v = get_sf_val('Cobertura temporal')
                if v:
                    add_label(doc, 'Cobertura temporal')
                    add_value(doc, v)

                v = get_sf_val('Última fecha de publicación')
                if v:
                    add_label(doc, 'Última fecha de publicación')
                    add_value(doc, v)

                v = get_sf_val('Versión de la base de datos')
                if v:
                    add_label(doc, 'Versión de la base de datos')
                    add_value(doc, v)

                v = get_sf_val('Formato de la base de datos')
                if v:
                    add_label(doc, 'Formato de la base de datos')
                    add_value(doc, v)

                vals = get_sf_vals('Fuente de información')
                if vals:
                    add_label(doc, 'Fuente de información')
                    if len(vals) > 1:
                        add_single_col_table(doc, vals)
                    else:
                        add_value(doc, vals[0])

                v = get_sf_val('Derechos')
                if v:
                    add_label(doc, 'Derechos')
                    add_value(doc, v)

                v = get_sf_val('Relaciones')
                if v:
                    add_label(doc, 'Relaciones')
                    add_value(doc, v)

                v = get_sf_val('Notas o comentarios')
                if v:
                    add_label(doc, 'Notas o comentarios sobre la base de datos')
                    add_value(doc, v)

    add_spacer(doc, 1)

    # ── MATERIALES DE REFERENCIA EXTERNOS ────────────────────────────────────
    sec_mat = find_top_sec(data, 'materiales')
    if sec_mat:
        e = find_elem(sec_mat, 'Descripción de los materiales')
        if e and has_content(e):
            add_chapter(doc, 'Materiales de Referencia Externos')
            add_section(doc, 'Descripción de los materiales de soporte conceptual y metodológico*')

            subfields = e.get('elemento', [])
            parent_entries = sorted(
                [c for c in e.get('contenido', []) if c.get('idPadreContenido') is None],
                key=lambda c: c.get('orden', 0)
            )

            def sf_val(parent_id, nombre_buscar):
                sf = next((s for s in subfields if nombre_buscar.lower() in s['nombre'].strip().lower()), None)
                if not sf:
                    return ''
                match = next((c for c in sf.get('contenido', []) if c.get('idPadreContenido') == parent_id), None)
                return fmt(sf, match['valor']) if match and match.get('valor') is not None else ''

            headers = ['Título del documento o material de referencia', 'URI o archivo de consulta',
                       'Tipo', 'Autor', 'Fecha', 'Descripción del contenido']
            rows = []
            for pe in parent_entries:
                pid = pe['idContenido']
                titulo = sf_val(pid, 'Título del documento')
                uri    = sf_val(pid, 'URI de consulta')
                tipo   = clean_cat(sf_val(pid, 'Tipo'))
                autor  = sf_val(pid, 'Autor')
                fecha  = sf_val(pid, 'Fecha')
                desc   = sf_val(pid, 'Descripción del contenido')
                row = [titulo, uri, tipo, autor, fecha, desc]
                if any(v for v in row):
                    rows.append(row)

            # Tabla con ancho total fijo 24 cm, columnas distribuidas uniformemente
            if rows:
                n_cols = len(headers)
                # 24 cm × 567 DXA/cm = 13,608 DXA total; cada columna = 2,268 DXA (~4 cm)
                total_dxa = 13608
                col_w = total_dxa // n_cols
                widths = [col_w] * n_cols
                widths[-1] = total_dxa - col_w * (n_cols - 1)

                table = doc.add_table(rows=1 + len(rows), cols=n_cols)
                table.style = 'Table Grid'
                set_table_width(table, total_dxa)

                # Desactivar autofit para que las columnas no se expandan por el contenido
                tbl = table._tbl
                tblPr = tbl.find(qn('w:tblPr'))
                tblLayout = tblPr.find(qn('w:tblLayout')) if tblPr is not None else None
                if tblLayout is None and tblPr is not None:
                    tblLayout = OxmlElement('w:tblLayout')
                    tblPr.append(tblLayout)
                if tblLayout is not None:
                    tblLayout.set(qn('w:type'), 'fixed')

                # Definir grid de columnas uniformes
                tblGrid = tbl.find(qn('w:tblGrid'))
                if tblGrid is None:
                    tblGrid = OxmlElement('w:tblGrid')
                    tbl.insert(1, tblGrid)
                else:
                    for gc in list(tblGrid):
                        tblGrid.remove(gc)
                for w in widths:
                    gc = OxmlElement('w:gridCol')
                    gc.set(qn('w:w'), str(w))
                    tblGrid.append(gc)

                for j, h in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    # Forzar ancho en cada celda
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.insert(0, tcW)
                    tcW.set(qn('w:w'), str(widths[j]))
                    tcW.set(qn('w:type'), 'dxa')
                    write_cell(cell, h, header=True)

                for i, row in enumerate(rows):
                    for j, val in enumerate(row):
                        cell = table.rows[i + 1].cells[j]
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is None:
                            tcW = OxmlElement('w:tcW')
                            tcPr.insert(0, tcW)
                        tcW.set(qn('w:w'), str(widths[j]))
                        tcW.set(qn('w:type'), 'dxa')
                        write_cell(cell, val or '')

                add_para(doc, '', before=0, after=4)

            # ── Nota del área encargada ──
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(12)
            r1 = p.add_run('*Nota del área encargada de la verificación de metadatos:')
            r1.font.name = FUENTE
            r1.font.size = Pt(10)
            r1.font.bold = True
            r1.font.color.rgb = GRIS
            r2 = p.add_run(' se informa que los Materiales de Referencia Externos cuentan con elementos de metadatos adicionales (Formato, Idioma, Colaboradores, Temas, Fuente, Cobertura, Derechos, Colaboradores, Editor) que no fueron incorporados en el presente formato, con el fin de concentrar el análisis en los metadatos referenciales que presentaban mayores áreas de oportunidad.')
            r2.font.name = FUENTE
            r2.font.size = Pt(10)
            r2.font.bold = False
            r2.font.color.rgb = GRIS

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(4)
            p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.line_spacing = Pt(12)
            r3 = p2.add_run('La totalidad de estos elementos se encuentra documentada en las fichas de metadatos (SIM) correspondientes a cada uno de los recursos incluidos en esta sección y podrá ser proporcionada de manera inmediata en caso de requerirse.')
            r3.font.name = FUENTE
            r3.font.size = Pt(10)
            r3.font.bold = False
            r3.font.color.rgb = GRIS

    return doc


# ─── Entry point ─────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) != 3:
        print('Uso: python sim_json_to_word.py input.json output.docx')
        sys.exit(1)
    input_path  = sys.argv[1]
    output_path = sys.argv[2]
    with open(input_path, encoding='utf-8-sig') as f:
        data = json.load(f)
    doc = build_document(data)
    doc.save(output_path)
    print(f'✓ Documento generado: {output_path}')


if __name__ == '__main__':
    main()
